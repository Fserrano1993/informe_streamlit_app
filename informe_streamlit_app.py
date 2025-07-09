#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Informe Generator App (FreeSimpleGUI)

Requisitos:
    pip install FreeSimpleGUI python-docx pdfplumber pdf2image pillow pandas openpyxl
"""

import os
import re
import json
import tempfile
from pathlib import Path

import pandas as pd
import FreeSimpleGUI as sg
import pdfplumber
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

CFG_FILE = "config_informes.json"
IMG_PH = "{{IMG_CATASTRO}}"
POPPLER = None

def load_cfg():
    if os.path.exists(CFG_FILE):
        return json.load(open(CFG_FILE, 'r', encoding='utf-8'))
    return {'base':'', 'jur':'', 'xls':''}

def save_cfg(cfg):
    with open(CFG_FILE, 'w', encoding='utf-8') as f:
        json.dump(cfg, f, indent=2)

def normaliza_modelo(modelo):
    if not modelo:
        return ''
    m = modelo.upper().split()[0].split('ED.')[0]
    return re.sub(r'-GEN|-SXXI|-PCI|-CO|-TR', '', m)

def modelo_a_ramo(modelo, xls):
    if not (modelo and xls and os.path.exists(xls)):
        return ''
    clave = normaliza_modelo(modelo)
    try:
        df = pd.read_excel(xls, engine='openpyxl')
        row = df[df.iloc[:,0].str.startswith(clave, na=False)]
        return row.iloc[0,1] if not row.empty else ''
    except:
        return ''

REG_ENC = {
    "{{EXPEDIENTE}}": r"Expediente:\s*(\S+)",
    "{{FECHA_DE_OCURRENCIA}}": r"Fecha de Ocurrencia:\s*([0-9/]+)",
    "{{EFECTO}}": r"Efecto:\s*([0-9/]+)",
    "{{GARANTIA_AFECTADA}}": r"Garantia afectada:\s*(.+)",
    "{{FECHA_HORA_SERVICIO}}": r"<NI>([0-9]{2}-[0-9]{2}-[0-9]{4})",
    "{{ASEGURADO}}": r"Asegurado:\s*([^\r\n]+)",
    "{{TLF1}}": r"Tlf1\s*[:\-]?\s*([0-9]+)",
    "{{MODELO_CONDICIONES_GENERALES}}": r"MODELO CONDICIONES GENERALES:\s*([^\r\n]+)",
    "{{AGUA_CONTENIDO}}": r"AGUA CONTENIDO:\s*([0-9\.,]+)",
    "{{AGUA_CONTINENTE}}": r"AGUA CONTINENTE:\s*([0-9\.,]+)",
    "{{DIR_ENCARGO}}": r"Lugar:\s*([^\r\n]+)"
}
REG_CAT = {
    "{{DIR_CATASTRO}}": r"Lugar:\s*([^\r\n]+)",
    "{{CP_CATASTRO}}": r"(\d{5})\s+",
    "{{LOCALIDAD_CATASTRO}}": r"\d{5}\s+([A-ZÃÃÃÃÃÃÃ\s]+)\[",
    "{{PROVINCIA_CATASTRO}}": r"\[([A-ZÃÃÃÃÃÃÃ]{2,})\]",
    "{{USO_PRINCIPAL_CATASTRAL}}": r"Uso principal:\s*([^\r\n]+?)(?:\s*Superficie|$)",
    "{{SUPERFICIE_CONSTRUIDA_CATASTRAL}}": r"Superficie construida:\s*([0-9.]+)",
    "{{ANO_CONSTRUCCION_CATASTRAL}}": r"AÃ±o construcciÃ³n:\s*([0-9]{4})",
    "{{SUPERFICIE_ELEMENTOS_COMUNES}}": r"Elementos comunes[^0-9]*([0-9.]+)",
    "{{PARTICIPACION_INMUEBLE}}": r"ParticipaciÃ³n del inmueble:\s*([0-9\.,]+ ?%)"
}

def parse_encargo(txt, cfg):
    rep = {k: '' for k in REG_ENC}
    for k, pat in REG_ENC.items():
        m = re.search(pat, txt, re.I)
        if m:
            rep[k] = m.group(1).strip()
    if rep.get("{{DIR_ENCARGO}}"):
        rep["{{DIR_CATASTRO}}"] = rep["{{DIR_ENCARGO}}"]
    f = rep.get("{{FECHA_DE_OCURRENCIA}}", "")
    parts = f.split('/')
    if len(parts)==3 and len(parts[2])==2:
        d,m,y = parts
        rep["{{FECHA_DE_OCURRENCIA}}"] = f"{d}/{m}/20{y}"
    rep["{{AGUA_CONTENIDO}}"] = rep.get("{{AGUA_CONTENIDO}}", "0")
    rep["{{AGUA_CONTINENTE}}"] = rep.get("{{AGUA_CONTINENTE}}", "0")
    rep["{{POLIZA_RAMO}}"] = modelo_a_ramo(rep.get("{{MODELO_CONDICIONES_GENERALES}}",""), cfg.get('xls',''))
    return rep

def parse_catastro(txt, rep):
    for k, pat in REG_CAT.items():
        if not rep.get(k):
            m = re.search(pat, txt, re.I)
            if m:
                rep[k] = m.group(1).strip()

def pdf_text(pdf):
    with pdfplumber.open(pdf) as p:
        return p.pages[0].extract_text() or ''

def pdf_png(pdf):
    img = convert_from_path(pdf, first_page=1, last_page=1, dpi=200, poppler_path=POPPLER)[0]
    tmp = Path(tempfile.gettempdir())/(Path(pdf).stem + '.png')
    img.save(tmp,'PNG')
    return str(tmp)

def replace_runs_placeholder(runs, ph, val):
    full = ''.join(r.text for r in runs)
    idx = full.find(ph)
    if idx < 0:
        return False
    start,end,pos = idx, idx+len(ph), 0
    for i,r in enumerate(runs):
        nxt = pos+len(r.text)
        if pos<=start<nxt:
            s_run,s_off = i, start-pos
        if pos<end<=nxt:
            e_run,e_off = i, end-pos
            break
        pos = nxt
    runs[s_run].text = runs[s_run].text[:s_off] + val + runs[e_run].text[e_off:]
    for j in range(s_run+1, e_run+1):
        runs[j].text = ''
    return True

def replace_paragraph(p, rep, img):
    if IMG_PH in p.text:
        p.clear()
        p.add_run().add_picture(img, width=Inches(6))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        for k,v in rep.items():
            if k in p.text:
                replace_runs_placeholder(p.runs, k, v)

def build_doc(tpl, img, out, rep):
    doc = Document(tpl)
    for para in doc.paragraphs:
        replace_paragraph(para, rep, img)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_paragraph(para, rep, img)
    doc.save(out)

def add_photo_report(doc, image_info):
    for i in range(0, len(image_info), 6):
        chunk = image_info[i:i+6]
        table = doc.add_table(rows=3, cols=2)
        table.autofit = True
        for idx,(imgp,cap) in enumerate(chunk):
            r = idx//2; c = idx%2
            cell = table.cell(r,c)
            p = cell.paragraphs[0]
            run = p.add_run()
            try:
                run.add_picture(imgp, width=Inches(2.5))
            except:
                p.add_run("[Error imagen]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_p = cell.add_paragraph(cap)
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cap_p.runs[0]
            cr.font.name = 'Calibri'
            cr.font.size = Pt(11)
        doc.add_page_break()

# ---- Interfaz ----
cfg = load_cfg()
sg.theme('SystemDefault')

tab_bib = [
    [sg.Text('Plantilla base'), sg.Input(cfg['base'], key='BASE', readonly=True, expand_x=True), sg.FileBrowse(file_types=(('Word','.docx'),))],
    [sg.Text('Plantilla jur.'), sg.Input(cfg['jur'], key='JUR', readonly=True, expand_x=True), sg.FileBrowse(file_types=(('Word','.docx'),))],
    [sg.Text('Excel ramos'),     sg.Input(cfg['xls'], key='XLS', readonly=True, expand_x=True), sg.FileBrowse(file_types=(('Excel','.xlsx'),))],
    [sg.Button('Guardar', key='SAVE')]
]
tab_inf = [
    [sg.Text('Encargo')],
    [sg.Radio('Adjuntar .txt','SRC', key='F', default=True, enable_events=True),
     sg.Radio('Pegar texto','SRC',  key='P', enable_events=True)],
    [sg.Input(key='ENC',expand_x=True), sg.FileBrowse(key='ENCBROWSE', file_types=(('Texto','.txt'),))],
    [sg.Multiline('', key='TXT', visible=False, disabled=True, size=(80,8))],
    [sg.Text('Catastro')],
    [sg.Input(key='IMG',expand_x=True), sg.FileBrowse(file_types=(('PDF/IMG','*.pdf;*.png;*.jpg'),))],
    [sg.Text('Fotos (carpeta)'), sg.Input(key='FOTOS', expand_x=True), sg.FolderBrowse()],
    [sg.Button('Generar', key='RUN'), sg.Button('Limpiar', key='CLR')]
]
layout = [[sg.TabGroup([[sg.Tab('Biblioteca',tab_bib), sg.Tab('Informe',tab_inf)]])]]
win = sg.Window('Generador de Informes', layout, finalize=True)

pie_opts = [
    "DaÃ±os en continente","DaÃ±os en contenido","Causa de los daÃ±os",
    "Vista general de la estancia","Vista general del riesgo","DaÃ±os reclamados"
]

while True:
    ev, val = win.read()
    if ev in (sg.WINDOW_CLOSED, None):
        break
    if ev in ('F','P'):
        pe = (ev=='P')
        win['ENC'].update(disabled=pe)
        win['ENCBROWSE'].update(disabled=pe)
        win['TXT'].update(visible=pe, disabled=not pe)
    if ev=='SAVE':
        cfg['base'],cfg['jur'],cfg['xls'] = val['BASE'],val['JUR'],val['XLS']
        save_cfg(cfg)
        sg.popup_ok('ConfiguraciÃ³n guardada')
    if ev=='CLR':
        for k in ('ENC','ENCBROWSE','TXT','IMG','FOTOS'):
            win[k].update('')
        win['F'].update(True); win['P'].update(False)
        win['ENC'].update(disabled=False); win['ENCBROWSE'].update(disabled=False)
        win['TXT'].update(visible=False, disabled=True)
    if ev=='RUN':
        try:
            texto = val['TXT'] if val['P'] else open(val['ENC'],'r',encoding='utf-8').read()
            rep = parse_encargo(texto, cfg)
            if val['IMG'].lower().endswith('.pdf'):
                ctxt = pdf_text(val['IMG'])
                parse_catastro(ctxt, rep)
                imgf = pdf_png(val['IMG'])
            else:
                imgf = val['IMG']
            tpl = cfg['jur'] if 'JURIDICA' in rep.get('{{GARANTIA_AFECTADA}}','').upper() else cfg['base']
            if not tpl:
                sg.popup_error('Error: Plantilla no definida')
                continue
            save_path = sg.popup_get_file(
                'Guardar informe como', save_as=True,
                default_path=f"{rep.get('{{EXPEDIENTE}}','SIN_EXP')}.docx",
                file_types=(('Word','.docx'),)
            )
            if not save_path:
                continue
            build_doc(tpl, imgf, save_path, rep)

            if val['FOTOS'] and os.path.isdir(val['FOTOS']):
                files = sorted(f for f in os.listdir(val['FOTOS']) if f.lower().endswith(('.png','.jpg','.jpeg')))
                tmpd = tempfile.mkdtemp()
                thumbs = []
                for i, fn in enumerate(files):
                    im = Image.open(os.path.join(val['FOTOS'], fn))
                    im.thumbnail((100,100))
                    tp = os.path.join(tmpd, f"thumb_{i}.png")
                    im.save(tp)
                    thumbs.append((os.path.join(val['FOTOS'], fn), tp))
                # Ventana pies
                layout2 = []
                for r in range(3):
                    row = []
                    for c in range(2):
                        idx = r*2+c
                        if idx < len(thumbs):
                            nm = os.path.basename(thumbs[idx][0])
                            row.append(sg.Column([
                                [sg.Image(thumbs[idx][1])],
                                [sg.Text(nm, size=(15,1))],
                                [sg.Combo(pie_opts, default_value=pie_opts[0], key=f"pie_{idx}")]
                            ], pad=(5,5)))
                        else:
                            row.append(sg.Column([[]], pad=(5,5)))
                    layout2.append(row)
                layout2.append([sg.Button('Aceptar'), sg.Button('Cancelar')])
                win2 = sg.Window('Pies de foto', layout2, modal=True)
                e2, v2 = win2.read()
                win2.close()
                if e2 == 'Aceptar':
                    sels = [v2[f"pie_{i}"] for i in range(len(thumbs))]
                    doc = Document(save_path)
                    add_photo_report(doc, list(zip([p[0] for p in thumbs], sels)))
                    doc.save(save_path)

            os.startfile(save_path)
        except Exception as e:
            sg.popup_error(f"Error al generar informe:\n{e}")

win.close()